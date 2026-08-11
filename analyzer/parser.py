import io
import re

import pdfplumber
import docx


class ResumeParseError(Exception):
    pass


def extract_text(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "pdf":
            return _extract_pdf(content)
        if ext in ("docx",):
            return _extract_docx(content)
        if ext == "txt":
            return content.decode("utf-8", errors="replace")
        if ext in ("doc",):
            raise ResumeParseError("Legacy .doc files are not supported. Please save the file as .docx or PDF and retry.")
        raise ResumeParseError(f"Unsupported file type: .{ext}. Please upload a PDF, DOCX or TXT file.")
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError(f"Failed to read file: {exc}") from exc


def _extract_pdf(content: bytes) -> str:
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages)


def _extract_docx(content: bytes) -> str:
    doc = docx.Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()
